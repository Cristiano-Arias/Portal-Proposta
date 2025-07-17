#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from flask import Flask, request, jsonify, send_from_directory, Response
from flask_cors import CORS
from datetime import datetime
import json
import os
import uuid
import logging
from io import BytesIO
import base64

# Imports para geração de documentos
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx não instalado. Funcionalidade Word limitada.")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("openpyxl não instalado. Funcionalidade Excel limitada.")

# Configuração do Flask
app = Flask(__name__, static_folder='static', static_url_path='')
app.config['JSON_AS_ASCII'] = False
app.config['JSONIFY_MIMETYPE'] = 'application/json; charset=utf-8'
CORS(app, origins=["*"], allow_headers=["Content-Type"], methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"])

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('propostas.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Diretórios
PROPOSTAS_DIR = 'propostas'
UPLOADS_DIR = 'uploads'
for directory in [PROPOSTAS_DIR, UPLOADS_DIR]:
    if not os.path.exists(directory):
        os.makedirs(directory)

# Base de dados em memória
propostas_db = {}
processos_db = {}

def salvar_proposta_arquivo(proposta_id, proposta_data):
    """Salva a proposta em arquivo JSON"""
    try:
        filename = os.path.join(PROPOSTAS_DIR, f'proposta_{proposta_id}.json')
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(proposta_data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        logger.error(f"Erro ao salvar proposta: {str(e)}")
        return False

def carregar_propostas():
    """Carrega todas as propostas do diretório"""
    try:
        for filename in os.listdir(PROPOSTAS_DIR):
            if filename.endswith('.json'):
                filepath = os.path.join(PROPOSTAS_DIR, filename)
                with open(filepath, 'r', encoding='utf-8') as f:
                    proposta = json.load(f)
                    propostas_db[proposta['id']] = proposta
    except Exception as e:
        logger.error(f"Erro ao carregar propostas: {str(e)}")

def calcular_totais_comerciais(dados):
    """Calcula totais da proposta comercial"""
    try:
        comercial = dados.get('comercial', {})
        
        # Converter valores string para float
        def str_to_float(value):
            if isinstance(value, str):
                return float(value.replace('.', '').replace(',', '.').replace('R$', '').strip())
            return float(value or 0)
        
        # Separar os totais
        total_servicos = str_to_float(comercial.get('totalServicos', 0))  # Total da Tabela de Serviços
        mao_obra = str_to_float(comercial.get('totalMaoObra', 0))
        materiais = str_to_float(comercial.get('totalMateriais', 0))
        equipamentos = str_to_float(comercial.get('totalEquipamentos', 0))
        
        # Custo direto = MO + Mat + Equip (SEM incluir serviços)
        custo_direto = mao_obra + materiais + equipamentos
        
        # BDI sobre o custo direto
        bdi_percentual = float(comercial.get('bdiPercentual', 0))
        bdi_valor = custo_direto * (bdi_percentual / 100)
        
        # Valor total = Custo Direto + BDI
        valor_total = custo_direto + bdi_valor
        
        # Formatar valores para exibição
        def format_currency(value):
            return f"R$ {value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        
        return {
            'servicos': format_currency(total_servicos),
            'mao_obra': format_currency(mao_obra),
            'materiais': format_currency(materiais),
            'equipamentos': format_currency(equipamentos),
            'custo_direto': format_currency(custo_direto),
            'bdi_percentual': f"{bdi_percentual:.1f}",
            'bdi_valor': format_currency(bdi_valor),
            'valor_total': format_currency(valor_total),
            'valor_total_num': valor_total
        }
    except Exception as e:
        logger.error(f"Erro ao calcular totais: {str(e)}")
        return {
            'servicos': 'R$ 0,00',
            'mao_obra': 'R$ 0,00',
            'materiais': 'R$ 0,00',
            'equipamentos': 'R$ 0,00',
            'custo_direto': 'R$ 0,00',
            'bdi_percentual': '0',
            'bdi_valor': 'R$ 0,00',
            'valor_total': 'R$ 0,00',
            'valor_total_num': 0
        }

def gerar_word_proposta(dados, protocolo):
    """Gera documento Word com a proposta técnica completa"""
    if not DOCX_AVAILABLE:
        return None
        
    try:
        doc = Document()
        
        # Configurar estilos
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Título Principal
        titulo = doc.add_heading('PROPOSTA TÉCNICA E COMERCIAL', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Protocolo
        p = doc.add_paragraph()
        p.add_run(f'Protocolo: {protocolo}').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data
        doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y")}')
        doc.add_paragraph()
        
        # 1. DADOS DA EMPRESA
        doc.add_heading('1. DADOS DA EMPRESA', level=1)
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Light Grid Accent 1'
        
        dados_empresa = [
            ('Razão Social:', dados.get('dados', {}).get('razaoSocial', '')),
            ('CNPJ:', dados.get('dados', {}).get('cnpj', '')),
            ('Endereço:', dados.get('dados', {}).get('endereco', '')),
            ('Cidade:', dados.get('dados', {}).get('cidade', '')),
            ('Telefone:', dados.get('dados', {}).get('telefone', '')),
            ('E-mail:', dados.get('dados', {}).get('email', '')),
            ('Responsável Técnico:', dados.get('dados', {}).get('respTecnico', '')),
            ('CREA/CAU:', dados.get('dados', {}).get('crea', ''))
        ]
        
        for i, (label, value) in enumerate(dados_empresa):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value or 'Não informado'
            table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        
        # 2. OBJETO DA CONCORRÊNCIA
        doc.add_page_break()
        doc.add_heading('2. OBJETO DA CONCORRÊNCIA', level=1)
        doc.add_paragraph(dados.get('tecnica', {}).get('objetoConcorrencia', ''))
        
        # 3. ESCOPO DOS SERVIÇOS
        doc.add_heading('3. ESCOPO DOS SERVIÇOS', level=1)
        
        doc.add_heading('3.1 Serviços Inclusos no Escopo', level=2)
        escopo_inclusos = dados.get('tecnica', {}).get('escopoInclusos', '')
        if escopo_inclusos:
            for item in escopo_inclusos.split('\n'):
                if item.strip():
                    doc.add_paragraph(f'• {item.strip()}', style='List Bullet')
        
        doc.add_heading('3.2 Serviços NÃO Inclusos no Escopo', level=2)
        escopo_excluidos = dados.get('tecnica', {}).get('escopoExclusos', '')
        if escopo_excluidos:
            doc.add_paragraph(escopo_excluidos)
        
        # 4. METODOLOGIA E PLANO DE ATAQUE
        doc.add_heading('4. METODOLOGIA E PLANO DE ATAQUE', level=1)
        
        doc.add_heading('4.1 Metodologia de Execução', level=2)
        doc.add_paragraph(dados.get('tecnica', {}).get('metodologia', ''))
        
        doc.add_heading('4.2 Sequência de Execução', level=2)
        doc.add_paragraph(dados.get('tecnica', {}).get('sequenciaExecucao', ''))
        
        # 5. PRAZO E CRONOGRAMA
        doc.add_page_break()
        doc.add_heading('5. PRAZO E CRONOGRAMA', level=1)
        
        doc.add_heading('5.1 Prazos', level=2)
        table_prazos = doc.add_table(rows=4, cols=2)
        table_prazos.style = 'Light Grid Accent 1'
        
        prazos = [
            ('Prazo Total de Execução:', dados.get('tecnica', {}).get('prazoExecucao', '') + ' dias'),
            ('Início dos Serviços:', dados.get('tecnica', {}).get('inicioServicos', 'Após assinatura do contrato')),
            ('Prazo de Mobilização:', dados.get('tecnica', {}).get('prazoMobilizacao', '15 dias')),
            ('Instalação do Canteiro:', dados.get('tecnica', {}).get('instalacaoCanteiro', '10 dias'))
        ]
        
        for i, (label, value) in enumerate(prazos):
            table_prazos.cell(i, 0).text = label
            table_prazos.cell(i, 1).text = value
            table_prazos.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        
        # Cronograma detalhado
        doc.add_heading('5.2 Cronograma Detalhado', level=2)
        cronograma = dados.get('tecnica', {}).get('cronograma', [])
        if cronograma and len(cronograma) > 0:
            table_crono = doc.add_table(rows=len(cronograma)+1, cols=4)
            table_crono.style = 'Light Grid Accent 1'
            
            headers = ['Atividade', 'Duração', 'Início', 'Fim']
            for i, header in enumerate(headers):
                cell = table_crono.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            for i, row in enumerate(cronograma):
                for j in range(min(4, len(row))):
                    table_crono.cell(i+1, j).text = str(row[j]) if row[j] else ''
        
        # 6. EQUIPE TÉCNICA
        doc.add_heading('6. EQUIPE TÉCNICA', level=1)
        equipe = dados.get('tecnica', {}).get('equipe', [])
        if equipe and len(equipe) > 0:
            table_equipe = doc.add_table(rows=len(equipe)+1, cols=5)
            table_equipe.style = 'Light Grid Accent 1'
            
            headers = ['Função', 'Nome', 'Formação', 'Experiência', 'Registro']
            for i, header in enumerate(headers):
                cell = table_equipe.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            for i, row in enumerate(equipe):
                for j in range(min(5, len(row))):
                    table_equipe.cell(i+1, j).text = str(row[j]) if row[j] else ''
        
        # 7. RECURSOS NECESSÁRIOS
        doc.add_page_break()
        doc.add_heading('7. RECURSOS NECESSÁRIOS', level=1)
        
        # Lista de Materiais
        doc.add_heading('7.1 Lista de Materiais', level=2)
        materiais = dados.get('tecnica', {}).get('materiais', [])
        if materiais and len(materiais) > 0:
            table_mat = doc.add_table(rows=len(materiais)+1, cols=4)
            table_mat.style = 'Light Grid Accent 1'
            
            headers = ['Material', 'Especificação', 'Unidade', 'Quantidade']
            for i, header in enumerate(headers):
                cell = table_mat.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            for i, row in enumerate(materiais):
                for j in range(min(4, len(row))):
                    table_mat.cell(i+1, j).text = str(row[j]) if row[j] else ''
        
        # Lista de Equipamentos
        doc.add_heading('7.2 Lista de Equipamentos', level=2)
        equipamentos = dados.get('tecnica', {}).get('equipamentos', [])
        if equipamentos and len(equipamentos) > 0:
            table_equip = doc.add_table(rows=len(equipamentos)+1, cols=4)
            table_equip.style = 'Light Grid Accent 1'
            
            headers = ['Equipamento', 'Especificação', 'Unidade', 'Quantidade']
            for i, header in enumerate(headers):
                cell = table_equip.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            for i, row in enumerate(equipamentos):
                for j in range(min(4, len(row))):
                    table_equip.cell(i+1, j).text = str(row[j]) if row[j] else ''
        
        # 8. ESTRUTURA DE CANTEIRO E OBRIGAÇÕES
        doc.add_heading('8. ESTRUTURA DE CANTEIRO E OBRIGAÇÕES', level=1)
        
        doc.add_heading('8.1 Estrutura do Canteiro', level=2)
        doc.add_paragraph(dados.get('tecnica', {}).get('estruturaCanteiro', ''))
        
        doc.add_heading('8.2 Obrigações da Contratada', level=2)
        doc.add_paragraph(dados.get('tecnica', {}).get('obrigacoesContratada', ''))
        
        doc.add_heading('8.3 Obrigações da Contratante', level=2)
        doc.add_paragraph(dados.get('tecnica', {}).get('obrigacoesContratante', ''))
        
        # 9. GARANTIAS
        doc.add_heading('9. GARANTIAS OFERECIDAS', level=1)
        doc.add_paragraph(dados.get('tecnica', {}).get('garantias', ''))
        
        # 10. EXPERIÊNCIA DA EMPRESA
        doc.add_heading('10. EXPERIÊNCIA DA EMPRESA', level=1)
        doc.add_paragraph(dados.get('tecnica', {}).get('experienciaEmpresa', ''))
        
        # 11. ATESTADOS TÉCNICOS
        doc.add_heading('11. ATESTADOS TÉCNICOS', level=1)
        doc.add_paragraph(dados.get('tecnica', {}).get('atestadosObras', ''))
        
        # 12. CONDIÇÕES COMERCIAIS
        doc.add_page_break()
        doc.add_heading('12. CONDIÇÕES COMERCIAIS', level=1)
        
        # Tabela resumo financeiro
        totais = calcular_totais_comerciais(dados)
        table_financeiro = doc.add_table(rows=8, cols=2)
        table_financeiro.style = 'Light Grid Accent 1'
        
        valores_financeiros = [
            ('Mão de Obra:', totais['mao_obra']),
            ('Materiais:', totais['materiais']),
            ('Equipamentos:', totais['equipamentos']),
            ('Custo Direto:', totais['custo_direto']),
            (f'BDI ({totais["bdi_percentual"]}%):', totais['bdi_valor']),
            ('VALOR TOTAL:', totais['valor_total']),
            ('Prazo de Execução:', dados.get('resumo', {}).get('prazoExecucao', '')),
            ('Validade da Proposta:', dados.get('comercial', {}).get('validadeProposta', '60 dias'))
        ]
        
        for i, (label, value) in enumerate(valores_financeiros):
            table_financeiro.cell(i, 0).text = label
            table_financeiro.cell(i, 1).text = str(value)
            table_financeiro.cell(i, 0).paragraphs[0].runs[0].font.bold = True
            if i == 5:  # Valor total
                table_financeiro.cell(i, 1).paragraphs[0].runs[0].font.bold = True
        
        # 13. CONDIÇÕES
        doc.add_heading('13. CONDIÇÕES', level=1)
        condicoes = dados.get('tecnica', {}).get('condicoesPremissas', '')
        if condicoes:
            doc.add_paragraph(condicoes)
        
        # 14. OBSERVAÇÕES FINAIS
        doc.add_heading('14. OBSERVAÇÕES FINAIS', level=1)
        doc.add_paragraph(dados.get('tecnica', {}).get('observacoesFinais', ''))
        
        # Assinatura
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('_' * 50)
        doc.add_paragraph(dados.get('dados', {}).get('respTecnico', ''))
        doc.add_paragraph(f'CREA/CAU: {dados.get("dados", {}).get("crea", "")}')
        doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y")}')
        
        # Salvar em memória
        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        
        return word_buffer
    except Exception as e:
        logger.error(f"Erro ao gerar Word: {str(e)}")
        return None

def gerar_excel_proposta(dados, protocolo):
    """Gera planilha Excel da proposta"""
    if not EXCEL_AVAILABLE:
        return None
        
    try:
        wb = Workbook()
        
        # Estilos
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        # Aba Resumo
        ws = wb.active
        ws.title = "Resumo"
        
        ws['A1'] = f"PROPOSTA {protocolo}"
        ws['A1'].font = Font(bold=True, size=16)
        
        # Dados da empresa
        ws['A3'] = "DADOS DA EMPRESA"
        ws['A3'].font = Font(bold=True)
        ws['A4'] = "Razão Social:"
        ws['B4'] = dados.get('dados', {}).get('razaoSocial', '')
        ws['A5'] = "CNPJ:"
        ws['B5'] = dados.get('dados', {}).get('cnpj', '')
        
        # Valores
        ws['A7'] = "RESUMO FINANCEIRO"
        ws['A7'].font = Font(bold=True)
        
        totais = calcular_totais_comerciais(dados)
        ws['A8'] = "Valor Total:"
        ws['B8'] = totais['valor_total']
        ws['B8'].font = Font(bold=True)
        
        # Salvar em memória
        excel_buffer = BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)
        
        return excel_buffer
    except Exception as e:
        logger.error(f"Erro ao gerar Excel: {str(e)}")
        return None

# Carrega propostas ao iniciar
carregar_propostas()

# ROTAS DA API

@app.route('/', methods=['GET'])
def home():
    """Serve o arquivo HTML principal ou info da API"""
    if os.path.exists(os.path.join('static', 'index.html')):
        return send_from_directory('static', 'index.html')
    else:
        return jsonify({
            "message": "Servidor de propostas funcionando - UTF-8 CORRIGIDO",
            "status": "online",
            "endpoints": [
                "/api/enviar-proposta",
                "/api/status",
                "/api/propostas/<id>",
                "/api/propostas/listar",
                "/api/processos/listar",
                "/api/verificar-cnpj",
                "/api/download/<tipo>/<id>"
            ]
        }), 200

@app.route('/api/status', methods=['GET'])
@app.route('/status', methods=['GET'])
def api_status():
    """Status do servidor"""
    return jsonify({
        "status": "online",
        "timestamp": datetime.now().isoformat(),
        "propostas_total": len(propostas_db),
        "processos_total": len(processos_db),
        "encoding": "UTF-8",
        "versao": "1.0.0",
        "mail_configured": os.environ.get('EMAIL_USER') is not None
    }), 200

@app.route('/api/enviar-proposta', methods=['POST', 'OPTIONS'])
@app.route('/enviar-proposta', methods=['POST', 'OPTIONS'])
def enviar_proposta():
    """Recebe e processa uma nova proposta"""
    if request.method == 'OPTIONS':
        return '', 200
        
    try:
        dados = request.get_json(force=True)
        logger.info(f"Recebendo proposta: {dados.get('protocolo')}")
        
        # Validação de CNPJ duplicado por processo
        cnpj = dados.get('dados', {}).get('cnpj', '').replace('.', '').replace('-', '').replace('/', '')
        processo = dados.get('processo', '')
        
        if cnpj and processo:
            # Verificar se já existe proposta deste CNPJ para este processo
            for prop_id, prop in propostas_db.items():
                cnpj_existente = prop.get('dados_empresa', {}).get('cnpj', '').replace('.', '').replace('-', '').replace('/', '')
                processo_existente = prop.get('processo', '')
                
                if cnpj_existente == cnpj and processo_existente == processo:
                    logger.warning(f"CNPJ duplicado: {cnpj} já tem proposta para processo {processo}")
                    return jsonify({
                        "sucesso": False,
                        "erro": "CNPJ já possui proposta para este processo",
                        "detalhes": f"Protocolo anterior: {prop.get('protocolo')}",
                        "data_anterior": prop.get('data_criacao')
                    }), 409  # Conflict
        
        # Gera ID único
        proposta_id = str(uuid.uuid4())
        
        # Estrutura completa da proposta
        proposta = {
            "id": proposta_id,
            "protocolo": dados.get('protocolo', f'PROP-{datetime.now().strftime("%Y%m%d%H%M%S")}'),
            "data_criacao": datetime.now().isoformat(),
            "status": "recebida",
            "dados_empresa": dados.get('dados', {}),
            "proposta_tecnica": dados.get('tecnica', {}),
            "proposta_comercial": dados.get('comercial', {}),
            "resumo": dados.get('resumo', {}),
            "processo": dados.get('processo', ''),
            "anexos": []
        }
        
        # Calcular totais
        totais = calcular_totais_comerciais(dados)
        proposta['totais_calculados'] = totais
        
        # Gerar documentos
        anexos = []
        
        # Gerar Word
        word_buffer = gerar_word_proposta(dados, proposta['protocolo'])
        if word_buffer:
            word_filename = f"proposta_{proposta['protocolo']}.docx"
            word_path = os.path.join(UPLOADS_DIR, word_filename)
            with open(word_path, 'wb') as f:
                f.write(word_buffer.getvalue())
            anexos.append({
                "tipo": "word",
                "nome": word_filename,
                "caminho": word_path
            })
        
        # Gerar Excel
        excel_buffer = gerar_excel_proposta(dados, proposta['protocolo'])
        if excel_buffer:
            excel_filename = f"proposta_{proposta['protocolo']}.xlsx"
            excel_path = os.path.join(UPLOADS_DIR, excel_filename)
            with open(excel_path, 'wb') as f:
                f.write(excel_buffer.getvalue())
            anexos.append({
                "tipo": "excel",
                "nome": excel_filename,
                "caminho": excel_path
            })
        
        proposta['anexos'] = anexos
        
        # Salvar proposta
        propostas_db[proposta_id] = proposta
        salvar_proposta_arquivo(proposta_id, proposta)
        
        return jsonify({
            "sucesso": True,
            "mensagem": "Proposta enviada com sucesso",
            "proposta_id": proposta_id,
            "protocolo": proposta['protocolo'],
            "anexos": [a['nome'] for a in anexos]
        }), 201
        
    except Exception as e:
        logger.error(f"Erro ao processar proposta: {str(e)}")
        return jsonify({
            "sucesso": False,
            "erro": "Erro ao processar proposta",
            "detalhes": str(e)
        }), 500

@app.route('/api/verificar-cnpj', methods=['POST'])
def verificar_cnpj():
    """Verifica se CNPJ já tem proposta para o processo"""
    try:
        data = request.get_json()
        cnpj = data.get('cnpj', '').replace('.', '').replace('-', '').replace('/', '')
        processo = data.get('processo', '')
        
        if not cnpj or not processo:
            return jsonify({"erro": "CNPJ e processo são obrigatórios"}), 400
        
        # Verificar duplicação
        for prop_id, prop in propostas_db.items():
            cnpj_existente = prop.get('dados_empresa', {}).get('cnpj', '').replace('.', '').replace('-', '').replace('/', '')
            processo_existente = prop.get('processo', '')
            
            if cnpj_existente == cnpj and processo_existente == processo:
                return jsonify({
                    "duplicado": True,
                    "protocolo": prop.get('protocolo'),
                    "data": prop.get('data_criacao'),
                    "empresa": prop.get('dados_empresa', {}).get('razaoSocial', '')
                }), 200
        
        return jsonify({"duplicado": False}), 200
        
    except Exception as e:
        logger.error(f"Erro ao verificar CNPJ: {str(e)}")
        return jsonify({"erro": "Erro ao verificar CNPJ"}), 500

@app.route('/api/propostas/<proposta_id>', methods=['GET'])
def obter_proposta(proposta_id):
    """Obtém uma proposta específica"""
    if proposta_id in propostas_db:
        return jsonify(propostas_db[proposta_id]), 200
    else:
        return jsonify({"erro": "Proposta não encontrada"}), 404

@app.route('/api/propostas/listar', methods=['GET'])
def listar_propostas():
    """Lista todas as propostas com paginação"""
    try:
        pagina = int(request.args.get('pagina', 1))
        por_pagina = int(request.args.get('por_pagina', 10))
        filtro_processo = request.args.get('processo', '')
        
        # Filtrar propostas
        propostas_lista = list(propostas_db.values())
        
        if filtro_processo:
            propostas_lista = [p for p in propostas_lista if filtro_processo.lower() in p.get('processo', '').lower()]
        
        # Ordenar por data
        propostas_lista.sort(key=lambda x: x['data_criacao'], reverse=True)
        
        # Paginação
        inicio = (pagina - 1) * por_pagina
        fim = inicio + por_pagina
        propostas_pagina = propostas_lista[inicio:fim]
        
        # Preparar resposta simplificada
        propostas_resumo = []
        for p in propostas_pagina:
            propostas_resumo.append({
                "id": p['id'],
                "protocolo": p['protocolo'],
                "empresa": p['dados_empresa'].get('razaoSocial', 'Não informado'),
                "cnpj": p['dados_empresa'].get('cnpj', ''),
                "valor": p.get('totais_calculados', {}).get('valor_total', 'R$ 0,00'),
                "data": p['data_criacao'],
                "status": p['status'],
                "processo": p.get('processo', '')
            })
        
        return jsonify({
            "propostas": propostas_resumo,
            "total": len(propostas_lista),
            "pagina": pagina,
            "por_pagina": por_pagina,
            "total_paginas": (len(propostas_lista) + por_pagina - 1) // por_pagina
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao listar propostas: {str(e)}")
        return jsonify({"erro": "Erro ao listar propostas"}), 500

@app.route('/api/processos/listar', methods=['GET'])
def listar_processos():
    """Lista todos os processos únicos das propostas"""
    try:
        processos = set()
        for proposta in propostas_db.values():
            processo = proposta.get('processo', '')
            if processo:
                processos.add(processo)
        
        return jsonify({
            "processos": sorted(list(processos)),
            "total": len(processos)
        }), 200
    except Exception as e:
        logger.error(f"Erro ao listar processos: {str(e)}")
        return jsonify({"erro": "Erro ao listar processos"}), 500

@app.route('/api/download/<tipo>/<proposta_id>', methods=['GET'])
def download_arquivo(tipo, proposta_id):
    """Download de arquivos anexos"""
    try:
        if proposta_id not in propostas_db:
            return jsonify({"erro": "Proposta não encontrada"}), 404
        
        proposta = propostas_db[proposta_id]
        
        # Procurar arquivo do tipo solicitado
        for anexo in proposta.get('anexos', []):
            if anexo['tipo'] == tipo and os.path.exists(anexo['caminho']):
                return send_from_directory(
                    os.path.dirname(anexo['caminho']),
                    os.path.basename(anexo['caminho']),
                    as_attachment=True,
                    download_name=anexo['nome']
                )
        
        return jsonify({"erro": f"Arquivo {tipo} não encontrado"}), 404
        
    except Exception as e:
        logger.error(f"Erro ao fazer download: {str(e)}")
        return jsonify({"erro": "Erro ao processar download"}), 500

@app.route('/api/propostas/<proposta_id>', methods=['PUT'])
def atualizar_proposta(proposta_id):
    """Atualiza o status de uma proposta"""
    try:
        if proposta_id not in propostas_db:
            return jsonify({"erro": "Proposta não encontrada"}), 404
        
        data = request.get_json()
        
        # Atualizar campos permitidos
        campos_atualizaveis = ['status', 'observacoes']
        for campo in campos_atualizaveis:
            if campo in data:
                propostas_db[proposta_id][campo] = data[campo]
        
        propostas_db[proposta_id]['data_atualizacao'] = datetime.now().isoformat()
        
        # Salvar alterações
        if salvar_proposta_arquivo(proposta_id, propostas_db[proposta_id]):
            return jsonify({
                "mensagem": "Proposta atualizada com sucesso",
                "proposta": propostas_db[proposta_id]
            }), 200
        else:
            return jsonify({"erro": "Erro ao salvar alterações"}), 500
            
    except Exception as e:
        logger.error(f"Erro ao atualizar proposta: {str(e)}")
        return jsonify({"erro": "Erro ao atualizar proposta"}), 500

# Servir arquivos estáticos
@app.route('/<path:path>')
def serve_static(path):
    """Serve arquivos estáticos da pasta static"""
    if os.path.exists(os.path.join('static', path)):
        return send_from_directory('static', path)
    else:
        return jsonify({"erro": "Arquivo não encontrado"}), 404

@app.errorhandler(404)
def nao_encontrado(e):
    return jsonify({"erro": "Endpoint não encontrado"}), 404

@app.errorhandler(500)
def erro_interno(e):
    return jsonify({"erro": "Erro interno do servidor"}), 500

if __name__ == '__main__':
    logger.info("Iniciando servidor de propostas...")
    # Configuração para Render
    port = int(os.environ.get('PORT', 5000))
    app.run(
        host='0.0.0.0',
        port=port,
        debug=False
    )
